#!/usr/bin/env python
# -*- coding:utf-8 -*-
# @Time  : 2019/10/14 16:33
# @Author: qimeng
# @File  : pdf2word.py

import os
import re
import datetime
from html import unescape
from tika import parser
from configparser import ConfigParser
from docx import Document

starttime = datetime.datetime.now()

def html_to_plain_text ( html ):
    text = re.sub('<head.*?>.*?</head>', '', html, flags=re.M | re.S | re.I)
    text = re.sub(r'<a\s.*?>', ' HYPERLINK ', text, flags=re.M | re.S | re.I)
    text = re.sub('<.*?>', '', text, flags=re.M | re.S)
    text = re.sub(r'-$\n', '', text, flags=re.M | re.S)
    text = re.sub(r'\n(\n+)', '---<<<///qm', text, flags=re.M | re.S)
    text = re.sub(r'(\s*\n)+', '', text, flags=re.M | re.S)
    text = re.sub('---<<<///qm', '\n', text, flags=re.M | re.S)
    if flag3 == '1':
        op = input("\n请选择替换方式：\n\t1：匹配仅包含全大写单词和空格的一行（eg：THE MATING MIND ）；\n" +
                   "\t2：匹配包含全大写单词,数字和空格的一行（eg：12 THE MATING MIND）；\n" +
                   "\t3：匹配不包含小写字母的一行（eg：F&  #IRSTANCHORB&OOKSEDITION , APRIL 2001 ）;\n" +
                   "\t0：手动输入正则表达式。\n")
        '''
        sou = input("\n请输入正则表达式：0\n\t匹配仅包含全大写单词和空格的一行（eg：THE MATING MIND ）：^[A-Z|\\s]+$\\n\n\t"+
                    "匹配包含全大写单词,数字和空格的一行（eg：12 THE MATING MIND）：^[A-Z|\\s|0-9]+$\\n\n\t"+
                    "匹配不好含小写字母的一行（eg：F&  #IRSTANCHORB&OOKSEDITION , APRIL 2001 ）：^[^a-z]+$\\n\n")
        '''
        if op == '1':
            text = re.sub('^[A-Z|\s]+$\n', '\n', text, flags=re.M | re.S)
        if op == '2':
            text = re.sub('^[A-Z|\\s|0-9]+$\n', '\n', text, flags=re.M | re.S)
        if op == '3':
            text = re.sub('^[^a-z]+$\n', '\n', text, flags=re.M | re.S)
        if op == '0':
            sou = input("\n请输入正则表达式：\n")
            des = input("\n请输入要替换为的内容：\n")
            text = re.sub(sou, des, text, flags=re.M | re.S)
        # text = re.sub('^[A-Z|\s]+$\n', '\n', text, flags=re.M | re.S)1

    if flag == '1':
        text = re.sub(r"([A-Z][a-zA-Z]+)\s*$\n", r"\1!qm! \n", text, flags=re.M | re.S)
        text = re.sub(r"(\.|\?|\!|。|？|！)$\n", r"\1 \n", text, flags=re.M | re.S)
        text = re.sub(r"(?<!((\.|\?|\!|。|？|！)\s))$\n", "", text, flags=re.M | re.S)
        text = re.sub(r"!qm!", r"", text, flags=re.M | re.S)

    if flag == '2':
        text = re.sub(r"(\.|\?|\!|。|？|！)$\n", r"\1 $\n", text, flags=re.M | re.S)
        text = re.sub(r"(?<!((\.|\?|\!|。|？|！)\s))$\n", "", text, flags=re.M | re.S)

    if flag2 == '1':
        text = re.sub(r'\s([a-zA-Z])\s([a-zA-Z])\s', r'\1\2', text, flags=re.M | re.S)
    if flag3 == '2':
        op = input("\n请选择替换方式：\n\t1：匹配仅包含全大写单词和空格的一行（eg：THE MATING MIND ）；\n" +
                   "\t2：匹配包含全大写单词,数字和空格的一行（eg：12 THE MATING MIND）；\n" +
                   "\t3：匹配不包含小写字母的一行（eg：F&  #IRSTANCHORB&OOKSEDITION , APRIL 2001 ）;\n" +
                   "\t0：手动输入正则表达式。\n")
        if op == '1':
            text = re.sub('^[A-Z|\s]+$\n', '\n', text, flags=re.M | re.S)
        if op == '2':
            text = re.sub('^[A-Z|\\s|0-9]+$\n', '\n', text, flags=re.M | re.S)
        if op == '3':
            text = re.sub('^[^a-z]+$\n', '\n', text, flags=re.M | re.S)
        if op == '0':
            sou = input("\n请输入正则表达式：\n")
            des = input("\n请输入要替换为的内容：\n")
            text = re.sub(sou, des, text, flags=re.M | re.S)

    return unescape(text)


def save_text_to_word(content1, file_path):
    doc = Document()
    for line in content1.split('\n'):
        paragraph = doc.add_paragraph()
        paragraph.add_run(remove_control_characters(line))
    doc.save(file_path)


def remove_control_characters(content2):
        mpa = dict.fromkeys(range(32))
        return content2.translate(mpa)


config_parser = ConfigParser()
config_parser.read('config.cfg')
config = config_parser['default']

tasks = []

for file in os.listdir(config['pdf_folder']):
    extension_name = os.path.splitext(file)[1]
    if extension_name != '.pdf':
        continue
    print("\n检测到新的pdf！")
    flag = input("\n是否需要处理换行？（此处回车后若未响应请再按一次回车）\n"+
                 "\t1：表示保留结尾符号和大写开头单词结尾换行;\n"
                 "\t2：表示仅保留结尾符号结尾换行;\n\t0：表示不处理。\n")
    flag2 = input("\n是否需要处理空格问题？;\n\t1：表示处理空格问题;\n\t0：表示不处理。\n")
    flag3 = input("\n是否需要正则字符串处理？;\n\t1：表示在换行空格处理前处理;\n\t2：表示在换行空格处理后处理;\n\t0：表示不处理。\n")
    file_name = os.path.splitext(file)[0]
    pdf_file = config['pdf_folder'] + '/' + file
    if flag2 == '1':
        file_name = 'S+' + file_name
    if flag == '1':
        file_name = 'L1+' + file_name
    if flag == '2':
        file_name = 'L2+' + file_name
    if flag3 == '1':
        file_name = 'D+' + file_name

    word_file = config['word_folder'] + '/' + file_name + '.docx'
    print("\n文件名：",file_name,"\n源路径：",pdf_file,"\n目标路径：",word_file)
    parse_entire_pdf = parser.from_file(pdf_file, xmlContent=True)
    parse_entire_pdf = parse_entire_pdf['content']
    content = html_to_plain_text(parse_entire_pdf)
    # save_text_to_word(parse_entire_pdf, word_file)
    save_text_to_word(content, word_file)


'''
pdf_file = 'D:/pdf2word/pdf/图论模板.pdf'
word_file = 'D:/pdf2word/word/冷记忆.docx'


parse_entire_pdf = parser.from_file(pdf_file, xmlContent=True)
parse_entire_pdf = parse_entire_pdf['content']
content = html_to_plain_text(parse_entire_pdf)

save_text_to_word(content, word_file)

'''


endtime = datetime.datetime.now()

print ("\nCPU耗时：",(endtime - starttime).seconds,"s")

os.system("pause")